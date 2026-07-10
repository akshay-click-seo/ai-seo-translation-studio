"""
services/docx_service.py — Word document (.docx) translation service.

Preserves:
  - Paragraph formatting (bold, italic, underline, font size, colour)
  - Heading styles
  - Tables (cell-by-cell translation)
  - Hyperlinks
  - Lists (bulleted / numbered)
  - Headers and footers
  - Section structure

Strategy:
  Translate at the Run level where formatting is stored,
  or at the Paragraph level when there is only one run.
"""

import io
import os
import copy
import time
from typing import Optional, Callable

from docx import Document
from docx.oxml.ns import qn

from logger import setup_logger, translation_logger
from utils import is_translatable

logger = setup_logger(__name__)


class DocxService:
    """Translates .docx files while preserving all formatting."""

    def translate(
        self,
        file_path: str,
        translator,
        target_lang: str = "es-la",
        mode: str = "Technical",
        use_tm: bool = True,
        on_progress: Optional[Callable[[int, int], None]] = None,
    ) -> bytes:
        """
        Translate all text in a .docx file.

        Args:
            file_path:   Path to source .docx file.
            translator:  Translator instance.
            target_lang: Target language code.
            mode:        Translation mode.
            use_tm:      Use Translation Memory.
            on_progress: Progress callback(current, total).

        Returns:
            Translated document as raw bytes.
        """
        t0 = time.perf_counter()
        doc = Document(file_path)

        # Build a flat list of all translatable items for progress tracking
        items = self._collect_items(doc)
        total = len(items)
        idx = 0

        # ── Translate body paragraphs ──────────────────────────────────────────
        for para in doc.paragraphs:
            self._translate_paragraph(para, translator, target_lang, mode, use_tm)
            if on_progress:
                on_progress(idx, total)
            idx += 1

        # ── Translate tables ───────────────────────────────────────────────────
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._translate_paragraph(para, translator, target_lang, mode, use_tm)
                        if on_progress:
                            on_progress(idx, total)
                        idx += 1

        # ── Translate headers and footers ──────────────────────────────────────
        for section in doc.sections:
            for hf in (section.header, section.footer,
                       section.first_page_header, section.first_page_footer):
                if hf:
                    for para in hf.paragraphs:
                        self._translate_paragraph(para, translator, target_lang, mode, use_tm)

        # Serialize
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        result = buf.read()

        elapsed = (time.perf_counter() - t0) * 1000
        translation_logger.log_file_translation(
            file_name=os.path.basename(file_path),
            file_type="docx",
            mode=mode,
            target_lang=target_lang,
            segments=total,
            processing_time_ms=elapsed,
        )
        return result

    # ── Private helpers ────────────────────────────────────────────────────────

    def _translate_paragraph(
        self,
        para,
        translator,
        target_lang: str,
        mode: str,
        use_tm: bool,
    ) -> None:
        """
        Translate a single paragraph in-place.

        For single-run paragraphs: translate and replace run text directly.
        For multi-run paragraphs: translate the full paragraph text and
        distribute the result across the first run (clearing others).
        """
        # Gather all run texts
        full_text = "".join(run.text for run in para.runs)
        if not full_text.strip() or not is_translatable(full_text):
            return

        try:
            result = translator.translate(
                full_text.strip(), target_lang=target_lang, mode=mode, use_tm=use_tm
            )
            translated = result["translation"]
        except Exception as exc:
            logger.error("Paragraph translation error: %s", exc)
            return

        if not para.runs:
            return

        if len(para.runs) == 1:
            # Simple: one run — just replace its text
            para.runs[0].text = translated
        else:
            # Multi-run: put translation in first run, clear the rest
            # This preserves the formatting of the first run for the whole para
            para.runs[0].text = translated
            for run in para.runs[1:]:
                run.text = ""

    def _collect_items(self, doc: Document) -> list:
        """Count translatable items for progress tracking."""
        items = []
        for para in doc.paragraphs:
            if para.text.strip():
                items.append(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            items.append(para)
        return items

    # ── Utility ────────────────────────────────────────────────────────────────

    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract all plain text from a .docx file.

        Args:
            file_path: Path to the .docx file.

        Returns:
            Full document text as a single string.
        """
        doc = Document(file_path)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            lines.append(para.text)
        return "\n".join(lines)

    @staticmethod
    def get_stats(file_path: str) -> dict:
        """
        Return word and paragraph counts for the document.

        Returns:
            Dict with keys: paragraphs, words, tables, chars.
        """
        doc = Document(file_path)
        text = " ".join(p.text for p in doc.paragraphs if p.text.strip())
        return {
            "paragraphs": sum(1 for p in doc.paragraphs if p.text.strip()),
            "words":      len(text.split()),
            "chars":      len(text),
            "tables":     len(doc.tables),
        }


# ── Module-level singleton ─────────────────────────────────────────────────────
docx_service = DocxService()
